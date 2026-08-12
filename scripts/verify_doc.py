#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Verify docx structure integrity: top tables, drawings, media, nested tables,
per-cell char counts.

Usage:
  python verify_doc.py FILE.docx
"""
import argparse, zipfile, re, lxml.etree as ET
from docx import Document

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('docx')
    args = ap.parse_args()
    d = Document(args.docx)
    print('顶层表数:', len(d.tables))
    with zipfile.ZipFile(args.docx) as z:
        doc = z.read('word/document.xml').decode('utf-8', 'replace')
        media = [n for n in z.namelist() if n.startswith('word/media/')]
    print('drawing 元素:', len(re.findall(r'<w:drawing', doc)))
    print('媒体文件:', media)
    root = ET.fromstring(zipfile.ZipFile(args.docx).read('word/document.xml'))
    tops = [c for c in root.find(W + 'body') if c.tag == W + 'tbl']
    nested = sum(1 for t in tops for tbl in t.iter(W + 'tbl') if tbl is not t)
    print('嵌套表(额外):', nested)
    print('--- 单元格字数 ---')
    for ti, t in enumerate(d.tables):
        for ri, r in enumerate(t.rows):
            for ci, c in enumerate(r.cells):
                txt = c.text.strip()
                if txt:
                    print(f'  [{ti}.{ri}.{ci}] {len(txt):5d}')

if __name__ == '__main__':
    main()
