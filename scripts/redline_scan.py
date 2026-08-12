#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Scan a .docx for forbidden/red-line strings and report counts + context.

Usage:
  python redline_scan.py FILE.docx [--extra w1 w2 ...]
"""
import argparse, re
from docx import Document

DEFAULT_REDLINES = [
    '王蒙', '(56)', '双层独立检索', 'patenthub', '广东省洛仑兹技术股份有限公司',
    'CN121278873A', '工装系统层', '第一次审查意见通知书', '审查员下发',
    '同日申请实用新型',  # should only appear as negation
]

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('docx')
    ap.add_argument('--extra', nargs='*', default=[])
    args = ap.parse_args()
    d = Document(args.docx)
    full = '\n'.join(c.text for t in d.tables for r in t.rows for c in r.cells)
    red = DEFAULT_REDLINES + args.extra
    print(f'=== 红线扫描: {args.docx} ===')
    hit = False
    for k in red:
        n = full.count(k)
        if n:
            hit = True
            print(f'  X 命中 {k!r} x{n}')
    if not hit:
        print('  [OK] 全文档红线 0 命中')
    # 同日申请/一案两报 context check (must be negation)
    for m in re.finditer(r'.{0,20}(同日申请|一案两报).{0,20}', full):
        seg = m.group(0).replace('\n', ' ')
        print('    上下文(应是否定句):', seg)
    print(f'顶层表数: {len(d.tables)}')
    print('python-docx 解析: OK')

if __name__ == '__main__':
    main()
